import glob
from pprint import pprint
import re
import csv
from docx import Document # pip install python-docx

ROOT_PATH = "."
# ファイルの場所を指定
file_path = f"{ROOT_PATH}/data/j24_probrem_zenki/"
# 実行済
newfile_name = f"{ROOT_PATH}/data/j24_problems_zenki.csv"

file_names = glob.glob(f'{file_path}2024*.docx')

file_contents = []
pattern = re.compile(r'No.*.c')


with open(newfile_name, mode='w', newline='', encoding='utf-8') as file:
    writer = csv.writer(file)
    writer.writerow(['date','problem_id', 'problem_text'])

for file_name in file_names:
    # docxファイルを開く
    doc = Document(file_name)
    contents_list = []
    tables_list = []

    for paragraph in doc.paragraphs:
        contents_list.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tables_list.append(cell.text)

    # content_listを結合し，空白文字'\u3000'を削除
    if  bool(pattern.search(''.join(contents_list))) :
        problems_id = [content for content in contents_list if "課題番号" in content]
        problems_id = [match.group() for id in problems_id if (match := pattern.search(id)) ]

        problems_text = tables_list[-1*len(problems_id):]
        problems_text = [str(text).replace("\u3000", '') for text in problems_text]

        #print(file_name)
        #print(problems_id)
        #print(problems_text)

        if len(problems_id) == len(problems_text) != 0:
            with open(newfile_name, mode='a', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                for problem_id, problem_text in zip(problems_id, problems_text):
                    data = "j2pro"+str(problem_id[2:6].zfill(4)) # 月日を追加
                    writer.writerow([data, problem_id, problem_text])


        #print("-------------------")
